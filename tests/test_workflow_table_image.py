import os
import sys
import pandas as pd
import docx
from PIL import Image, ImageDraw
from unittest.mock import patch
import types

tests_dir = os.path.abspath(os.path.dirname(__file__))
sys.path.insert(0, tests_dir)
sys.path.insert(0, os.path.abspath(os.path.join(tests_dir, "..")))

dummy_st = types.ModuleType("sentence_transformers")
class DummyModel:
    def encode(self, *args, **kwargs):
        return [0.0]

dummy_st.SentenceTransformer = DummyModel
sys.modules.setdefault("sentence_transformers", dummy_st)

dummy_pymysql = types.ModuleType("pymysql")
dummy_pymysql.connect = lambda *a, **kw: None
sys.modules.setdefault("pymysql", dummy_pymysql)

from src.core.docx_parser.docx_process import ParserDocx
from src.core.save_to_cache.insert2mo import insertAY2mo
from src.core.save_to_cache.mo import MatrixOne
from debug_helpers import debug_read2docx, Timeout


class DummySplitter:
    def __init__(self, model=None):
        pass

    def split_passages(self, text):
        return [text]

class DummyEmbedding:
    def get_embedding(self, text):
        return [0.0]

def create_doc(path, img_path):
    doc = docx.Document()
    table1 = doc.add_table(rows=2, cols=2)
    table1.cell(0, 0).text = "A"
    table1.cell(0, 1).text = "B"
    table1.cell(1, 0).text = "1"
    table1.cell(1, 1).text = "2"
    doc.add_page_break()
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "A"
    table2.cell(0, 1).text = "B"
    table2.cell(1, 0).text = "3"
    table2.cell(1, 1).text = "4"
    doc.add_picture(str(img_path))
    doc.save(path)

def create_image(path):
    img = Image.new("RGB", (50, 20), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((0, 0), "data", fill="black")
    img.save(path)

def test_workflow_table_image(tmp_path):
    doc_path = tmp_path / "test.docx"
    img_path = tmp_path / "img.png"
    create_image(img_path)
    create_doc(doc_path, img_path)

    parser = ParserDocx()
    table_file = tmp_path / "t.xlsx"
    with patch("src.core.utils.ocr.image_to_text", return_value="data"), \
         patch("src.core.docx_parser.docx_process.ocr_image_to_text", return_value="data"), \
         patch.object(ParserDocx, "convert_rows_to_md", return_value=str(table_file)):
        with Timeout(10):
            parsed = debug_read2docx(parser, str(doc_path))
    pd.DataFrame({"A": [1]}).to_excel(table_file, index=False)
    tables = [r for r in parsed if r["type"] == "tables"]
    assert len(tables) == 1
    assert tables[0]["content"].endswith(".xlsx")
    for t in tables:
        t["table_name"] = "table"

    images = [r for r in parsed if r.get("style") == "image"]
    assert images

    inserted = []
    def fake_insert(self, table, columns, values):
        inserted.append((table, columns, values))

    with patch.object(MatrixOne, "_insert_to_table", new=fake_insert), \
         patch.object(MatrixOne, "__init__", lambda self: None), \
         patch("src.core.save_to_cache.insert2mo.table_message", return_value="summary"), \
         patch("src.core.save_to_cache.insert2mo.SemanticParagraphSplitter", DummySplitter):
        insertAY2mo("name", "file", tables, parsed, DummyEmbedding(), own=True)

    assert any("table" in vals for _, cols, vals in inserted)
