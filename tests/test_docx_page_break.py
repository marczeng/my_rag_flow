import os
import sys
import docx

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.core.docx_parser.docx_process import ParserDocx


def create_doc(path):
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.add_page_break()
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "A"
    table2.cell(0, 1).text = "B"
    table2.cell(1, 0).text = "3"
    table2.cell(1, 1).text = "4"
    doc.save(path)


def test_page_break_metadata(tmp_path):
    file_path = tmp_path / "pb.docx"
    create_doc(str(file_path))
    parser = ParserDocx()
    res = parser.read2docx(str(file_path))
    assert all("page" in r for r in res)
    tables = [r for r in res if r["type"] == "tables"]
    assert len(tables) == 1
    assert tables[0]["page"] == 1
