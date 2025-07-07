import pandas as pd
import docx
from src.core.docx_parser.docx_process import ParserDocx


def create_split_doc(path):
    doc = docx.Document()
    table1 = doc.add_table(rows=2, cols=2)
    table1.cell(0, 0).text = "A"
    table1.cell(0, 1).text = "B"
    table1.cell(1, 0).text = "1"
    table1.cell(1, 1).text = "2"

    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "A"
    table2.cell(0, 1).text = "B"
    table2.cell(1, 0).text = "3"
    table2.cell(1, 1).text = "4"
    doc.save(path)


def test_split_tables_combined(tmp_path):
    file_path = tmp_path / "split.docx"
    create_split_doc(file_path)
    parser = ParserDocx()
    result = parser.read2docx(str(file_path))
    tables = [r for r in result if r["type"] == "tables"]
    assert len(tables) == 1
    df = pd.read_excel(tables[0]["content"])
    assert list(df.columns) == ["A", "B"]
    assert df.shape[0] == 2
    assert str(df.iloc[0, 0]) == "1"
    assert str(df.iloc[1, 0]) == "3"
